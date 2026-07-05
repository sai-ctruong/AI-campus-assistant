"""Parse file Word (.docx) → text. Lấy cả đoạn văn lẫn nội dung bảng."""
from docx import Document

def parse_docx(file_path: str) -> str:
    doc = Document(file_path)
    parts: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Nội dung trong bảng (mỗi hàng 1 dòng, các ô cách nhau bằng " | ").
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)
